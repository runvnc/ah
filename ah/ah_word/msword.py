from docx import Document

async def replace_all(file_path, replacements, save_path=None, context=None):
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if old_text in inline[i].text:
                        inline[i].text = inline[i].text.replace(old_text, new_text)
    
    save_path = save_path or file_path
    doc.save(save_path)
    return f"Replacements made and document saved to {save_path}"

async def read_document(file_path, context=None):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)
